
import streamlit as st
import os
import re
import json
import random
from typing import List, Dict, Tuple
import pypdf
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from io import BytesIO
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table
from reportlab.lib.units import inch

# ==========================================
# CONFIGURATION & SETUP
# ==========================================

st.set_page_config(
    page_title="Smart Exam Generator",
    page_icon="📄",
    layout="wide"
)

# Custom CSS
st.markdown("""
<style>
    .main-header { font-size: 2.5rem; font-weight: bold; color: #1f1f1f; text-align: center; margin-bottom: 1rem; }
    .sub-header { font-size: 1.2rem; color: #555; text-align: center; margin-bottom: 2rem; }
    .exam-paper { background-color: #ffffff; padding: 30px; border: 1px solid #ddd; box-shadow: 0 4px 6px rgba(0,0,0,0.1); border-radius: 5px; margin-top: 20px; font-family: 'Times New Roman', Times, serif; }
    .section-header { background-color: #f0f0f0; padding: 10px; font-weight: bold; margin-top: 20px; border-left: 5px solid #333; }
    .question-item { margin-bottom: 15px; padding: 10px 0; border-bottom: 1px dotted #ccc; }
    .download-box { border: 1px solid #eee; padding: 15px; border-radius: 10px; margin-top: 10px; background-color: #fafafa; }
</style>
""", unsafe_allow_html=True)

# ==========================================
# SESSION STATE
# ==========================================

if 'step' not in st.session_state:
    st.session_state.step = 1
if 'full_text' not in st.session_state:
    st.session_state.full_text = ""
if 'chapters_data' not in st.session_state:
    st.session_state.chapters_data = []
if 'generated_exam' not in st.session_state:
    st.session_state.generated_exam = None

# ==========================================
# LOCAL LOGIC ENGINE (NO API)
# ==========================================

def analyze_structure_locally(text: str) -> List[Dict]:
    if not text or len(text.strip()) < 10:
        return [{"name": "Full Document", "content": text}]
        
    chapters = []
    try:
        pattern = r'(?:\n|^)(Chapter|Unit|Section|Part)\s+\d+[:\-.]?\s*(.*)'
        matches = list(re.finditer(pattern, text, re.IGNORECASE))
        
        if len(matches) > 1:
            for i, match in enumerate(matches):
                start = match.start()
                end = matches[i+1].start() if i < len(matches) - 1 else len(text)
                title = match.group(0).strip()
                content = text[start:end].strip()
                if len(content) > 20:
                    chapters.append({"name": title, "content": content})
        else:
            st.info("No clear chapter headings found. Document is split into 5 equal sections automatically.")
            chunk_size = len(text) // 5
            for i in range(5):
                start = i * chunk_size
                end = (i + 1) * chunk_size if i < 4 else len(text)
                content = text[start:end].strip()
                if len(content) > 50:
                    chapters.append({"name": f"Section {i+1}", "content": content})
    except Exception as e:
        st.warning(f"Could not analyze structure automatically. Using full text. Error: {e}")
        return [{"name": "Full Document", "content": text}]
                
    if not chapters:
        return [{"name": "Full Document", "content": text}]
    return chapters

def generate_mcq_question(sentence: str, all_sentences: list) -> Dict:
    """
    Attempts to convert a sentence into a standard MCQ (Who/What question)
    instead of a fill-in-the-blank.
    """
    words = sentence.split()
    if len(words) < 4:
        return None

    candidates = [w for w in words if len(w) > 4 and w.isalpha()]
    if not candidates:
        candidates = [w for w in words if len(w) > 4]
    if not candidates: 
        return None

    target_word = random.choice(candidates)
    target_index = words.index(target_word)
    
    # --- SMART CONVERSION LOGIC ---
    # Pattern: "Subject is Object" -> "What is Object?" or "What is Subject?"
    # Pattern: "Subject verb Object" -> "What does Subject verb?"
    
    verb_pattern = r'\b(is|are|was|were)\b'
    has_linking_verb = re.search(verb_pattern, sentence, re.IGNORECASE)
    
    question_text = ""
    
    try:
        if has_linking_verb:
            # Attempt to split sentence by the verb
            split_idx = -1
            for idx, w in enumerate(words):
                if w.lower() in ['is', 'are', 'was', 'were']:
                    split_idx = idx
                    break
            
            if split_idx != -1:
                subject = " ".join(words[:split_idx])
                predicate = " ".join(words[split_idx+1:])
                
                # If we masked the Subject (e.g. "Paris is...")
                if target_index < split_idx:
                    # Ask: "What is the predicate?"
                    # Take only first few words of predicate for brevity
                    predicate_snippet = " ".join(predicate.split()[:5])
                    question_text = f"What {words[split_idx]} {predicate_snippet}?"
                
                # If we masked the Object/Predicate (e.g. "...is the capital.")
                else:
                    # Ask: "What is the Subject?"
                    question_text = f"What {words[split_idx]} {subject}?"
                question_text = question_text.strip(" ?") + "?"
        
        # Fallback: If we couldn't make a 'What' question (or it was in the middle)
        # Use Sentence Completion format: "The capital of France is:"
        if not question_text:
            # Remove the target word
            temp_words = words.copy()
            temp_words[target_index] = ""
            # Clean up double spaces
            clean_sentence_part = " ".join(temp_words).replace("  ", " ").strip()
            # Add a colon to show it's incomplete
            question_text = f"Complete the sentence: {clean_sentence_part}:"
            
            # If the target word was the last one, it looks great: "The capital of France is:"
            # If it was in the middle, it looks okay: "The sky is very:"

    except Exception:
        # Ultimate Fallback: Fill in the blank (labeled as MCQ)
        masked_words = words.copy()
        masked_words[target_index] = "______"
        question_text = " ".join(masked_words)

    # Generate Distractors
    distractors = set()
    attempts = 0
    while len(distractors) < 3 and attempts < 20:
        attempts += 1
        rand_sent = random.choice(all_sentences)
        rand_words = [w for w in rand_sent.split() if len(w) > 4]
        if rand_words:
            d = random.choice(rand_words)
            if d.lower() != target_word.lower():
                distractors.add(d)
    
    # If we couldn't find enough distractors, create generic ones
    # (This prevents crashing on short texts)
    while len(distractors) < 3:
        distractors.add(f"Option {len(distractors)+1}")

    options = list(distractors) + [target_word]
    random.shuffle(options)
    
    option_map = {chr(65+i): opt for i, opt in enumerate(options)}
    correct_letter = [k for k, v in option_map.items() if v == target_word][0]
    
    return {
        "question": question_text,
        "options": [f"{k}. {v}" for k, v in option_map.items()],
        "correct": correct_letter
    }

def local_generate_exam(text_source: str, config: dict) -> Dict:
    if not text_source:
        return {"error": "The selected text is empty."}

    raw_sentences = re.split(r'(?<=[.!?])\s+', text_source)
    sentences = [s.strip() for s in raw_sentences if len(s.strip()) > 20 and len(s.strip()) < 300]
    
    if len(sentences) < 3:
        return {"error": "Text is too short to generate questions."}

    random.shuffle(sentences)
    
    mcqs = []
    num_mcqs = min(config['num_mcqs'], len(sentences))
    
    for i in range(num_mcqs):
        try:
            q_data = generate_mcq_question(sentences[i], sentences)
            if q_data:
                mcqs.append(q_data)
        except Exception:
            continue

    shorts = []
    num_short = min(config['num_short'], len(sentences) - num_mcqs)
    for i in range(num_mcqs, num_mcqs + num_short):
        try:
            sentence = sentences[i]
            clean_q = sentence.rstrip('.')
            shorts.append(f"Define or Explain: {clean_q}")
        except Exception:
            continue

    longs = []
    num_long = min(config['num_long'], len(sentences) - (num_mcqs + num_short))
    for i in range(num_mcqs + num_short, num_mcqs + num_short + num_long):
        try:
            sentence = sentences[i]
            clean_q = sentence.rstrip('.')
            longs.append(f"Write a detailed note on: {clean_q}")
        except Exception:
            continue

    if not mcqs and not shorts and not longs:
        return {"error": "Could not generate any valid questions."}

    return {
        "subject": "Generated Exam",
        "class_level": config['academic_level'],
        "time": "2 Hours",
        "total_marks": str((num_mcqs * 1) + (num_short * 5) + (num_long * 10)),
        "section_a_mcqs": mcqs,
        "section_b_short": shorts,
        "section_c_long": longs
    }

# ==========================================
# EXPORT FUNCTIONS
# ==========================================

def escape_pdf_text(text):
    if not text: return ""
    return text.replace('&', '&amp;').replace('<', '&lt;').replace('>', '&gt;')

def create_word_document(exam: Dict, show_answers: bool = False, is_answer_key_only: bool = False) -> BytesIO:
    try:
        doc = Document()
        title = "MCQ Answer Key" if is_answer_key_only else "Exam Paper"
        head = doc.add_heading(title, 0)
        head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p = doc.add_paragraph()
        p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        p.add_run(f"Subject: {exam.get('subject', 'N/A')}\nClass: {exam.get('class_level', 'N/A')}\nTime: {exam.get('time', '2 Hours')}\n")
        
        if exam.get('section_a_mcqs'):
            doc.add_heading("SECTION A – MCQs", level=1)
            for i, q in enumerate(exam['section_a_mcqs'], 1):
                if is_answer_key_only:
                    p = doc.add_paragraph(f"Q{i}) {q.get('correct', '')}")
                else:
                    p = doc.add_paragraph(style='List Number')
                    p.add_run(q.get('question', '')).bold = True
                    for opt in q.get('options', []):
                        doc.add_paragraph(opt, style='List Bullet')
                    if show_answers:
                         doc.add_paragraph(f"Answer: {q.get('correct', '')}").italic = True

        if not is_answer_key_only:
            if exam.get('section_b_short'):
                doc.add_heading("SECTION B – Short Questions", level=1)
                for i, q in enumerate(exam['section_b_short'], 1):
                    p = doc.add_paragraph(style='List Number')
                    p.add_run(q)
            if exam.get('section_c_long'):
                doc.add_heading("SECTION C – Long Questions", level=1)
                for i, q in enumerate(exam['section_c_long'], 1):
                    p = doc.add_paragraph(style='List Number')
                    p.add_run(q)
        
        buffer = BytesIO()
        doc.save(buffer)
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error(f"Error creating Word Document: {e}")
        return None

def create_pdf_document(exam: Dict, show_answers: bool = False, is_answer_key_only: bool = False) -> BytesIO:
    try:
        buffer = BytesIO()
        doc = SimpleDocTemplate(buffer, pagesize=letter, leftMargin=0.75*inch, rightMargin=0.75*inch)
        story = []
        styles = getSampleStyleSheet()
        style_normal = styles["BodyText"]
        style_normal.leading = 18
        style_header = ParagraphStyle('CustomHeader', parent=styles['Heading1'], alignment=1) 
        
        title_text = "MCQ Answer Key" if is_answer_key_only else "Examination Paper"
        story.append(Paragraph(title_text, style_header))
        info_text = f"<b>Subject:</b> {exam.get('subject', '')} | <b>Class:</b> {exam.get('class_level', '')}<br/>Time: {exam.get('time', '')} | Marks: {exam.get('total_marks', '')}"
        story.append(Paragraph(info_text, style_normal))
        story.append(Spacer(1, 12))

        if exam.get('section_a_mcqs'):
            story.append(Paragraph("SECTION A – MCQs", styles['Heading2']))
            for i, q in enumerate(exam['section_a_mcqs'], 1):
                if is_answer_key_only:
                    text = f"{i}. {q.get('correct', '')}"
                    story.append(Paragraph(text, style_normal))
                else:
                    safe_q_text = escape_pdf_text(q.get('question', ''))
                    q_text = f"<b>Q{i}.</b> {safe_q_text}"
                    story.append(Paragraph(q_text, style_normal))
                    
                    for opt in q.get('options', []):
                        safe_opt = escape_pdf_text(opt)
                        story.append(Paragraph(f"• {safe_opt}", style_normal))
                        
                    if show_answers:
                        story.append(Paragraph(f"<i>Correct Answer: {q.get('correct', '')}</i>", style_normal))
                    story.append(Spacer(1, 6))

        if not is_answer_key_only:
            if exam.get('section_b_short'):
                story.append(Paragraph("SECTION B – Short Questions", styles['Heading2']))
                for i, q in enumerate(exam['section_b_short'], 1):
                    story.append(Paragraph(f"<b>Q{i}.</b> {escape_pdf_text(q)}", style_normal))
                    story.append(Spacer(1, 6))
            if exam.get('section_c_long'):
                story.append(Paragraph("SECTION C – Long Questions", styles['Heading2']))
                for i, q in enumerate(exam['section_c_long'], 1):
                    story.append(Paragraph(f"<b>Q{i}.</b> {escape_pdf_text(q)}", style_normal))
                    story.append(Spacer(1, 6))

        doc.build(story)
        buffer.seek(0)
        return buffer
    except Exception as e:
        st.error("Failed to generate PDF. Please use Word export instead.")
        return None

# ==========================================
# FILE HELPERS
# ==========================================

def extract_text_from_file(uploaded_file) -> str:
    text = ""
    try:
        if not uploaded_file:
            raise ValueError("No file uploaded.")
            
        if uploaded_file.type == "application/pdf":
            reader = pypdf.PdfReader(uploaded_file)
            if len(reader.pages) == 0:
                raise ValueError("PDF file is empty.")
            for page in reader.pages:
                extracted = page.extract_text()
                if extracted: text += extracted + "\n"
                
        elif uploaded_file.type == "application/vnd.openxmlformats-officedocument.wordprocessingml.document":
            doc = Document(uploaded_file)
            for para in doc.paragraphs:
                text += para.text + "\n"
                
        elif uploaded_file.type == "text/plain":
            text = str(uploaded_file.read(), "utf-8")
        else:
            st.error("Unsupported file type. Please upload PDF, DOCX, or TXT.")
            return ""
            
    except Exception as e:
        st.error(f"Error reading file: {e}")
        return ""
        
    return re.sub(r'\s+', ' ', text).strip()

# ==========================================
# UI LOGIC
# ==========================================

def main():
    st.markdown('<div class="main-header">Smart Exam Generator</div>', unsafe_allow_html=True)
    st.markdown('<div class="sub-header">Assist students and teachers in generating exams automatically from uploaded syllabus</div>', unsafe_allow_html=True)

    with st.sidebar:
        st.header("⚙️ Configuration")
        st.divider()
        with st.expander("ℹ️ Project Info", expanded=False):
            st.markdown("**Project Name:** Smart Exam Generator")
            st.markdown("**Developed By:** Muhammad Mudasar")
            st.markdown("**Organization:** SP Brothers Software House")
            st.markdown("---")
            st.markdown("This project assists students and teachers in generating exams automatically from uploaded syllabus.")
        st.divider()
        st.info("Rule-Based Engine: Converts text sentences into standard MCQs and Sentence Completion questions.")

    if st.session_state.step == 1:
        st.header("Step 1: Content Ingestion")
        uploaded_file = st.file_uploader("Upload Syllabus / Book", type=['pdf', 'docx', 'txt'])
        if uploaded_file:
            with st.spinner("Analyzing structure locally..."):
                try:
                    st.session_state.full_text = extract_text_from_file(uploaded_file)
                    if st.session_state.full_text:
                        st.session_state.chapters_data = analyze_structure_locally(st.session_state.full_text)
                        st.session_state.step = 2
                        st.rerun()
                    else:
                        st.error("Could not extract text.")
                except Exception as e:
                    st.error(f"An unexpected error occurred: {e}")

    elif st.session_state.step == 2:
        st.header("Step 2: Select Chapters / Sections")
        if not st.session_state.chapters_data:
            st.error("No chapters found.")
            if st.button("Back"): st.session_state.step = 1; st.rerun()
        else:
            st.info(f"Found {len(st.session_state.chapters_data)} sections.")
            options = [c['name'] for c in st.session_state.chapters_data]
            selected_indices = st.multiselect("Choose sections to test:", options=options, default=options)
            
            if st.button("Continue"):
                if selected_indices:
                    selected_content_list = [
                        c['content'] for c in st.session_state.chapters_data 
                        if c['name'] in selected_indices
                    ]
                    st.session_state.selected_text = "\n".join(selected_content_list)
                    st.session_state.step = 3
                    st.rerun()
                else:
                    st.warning("Please select at least one section.")

    elif st.session_state.step == 3:
        st.header("Step 3: Exam Configuration")
        with st.form("config"):
            col1, col2 = st.columns(2)
            with col1:
                exam_type = st.selectbox("Exam Type", ["Both", "MCQs only", "Subjective only"])
                diff = st.selectbox("Difficulty (Visual)", ["Mixed", "Easy", "Medium", "Hard"])
                lvl = st.selectbox("Level", ["School", "College", "University"])
            with col2:
                mcqs = st.slider("MCQs", 0, 20, 5)
                short = st.slider("Short", 0, 10, 3)
                long = st.slider("Long", 0, 5, 2)
            
            submitted = st.form_submit_button("Generate Locally")
            if submitted:
                with st.spinner("Generating questions from text..."):
                    cfg = {"exam_type": exam_type, "difficulty": diff, "academic_level": lvl, "num_mcqs": mcqs, "num_short": short, "num_long": long}
                    try:
                        result = local_generate_exam(st.session_state.selected_text, cfg)
                        
                        if result and "error" not in result:
                            st.session_state.generated_exam = result
                            st.session_state.step = 4
                            st.rerun()
                        else:
                            st.error(result.get("error", "Failed to generate exam."))
                    except Exception as e:
                        st.error(f"An unexpected error occurred during generation: {e}")

    elif st.session_state.step == 4:
        exam = st.session_state.generated_exam
        if not exam:
            st.error("No exam data found.")
            if st.button("Back"): st.session_state.step = 1; st.rerun()
        else:
            st.header("Generated Exam Paper")
            
            with st.container():
                st.markdown(f"""
                <div class="exam-paper">
                    <div style="text-align: center; margin-bottom: 30px;">
                        <h2>{exam.get('subject', 'Subject')}</h2>
                        <p><strong>Class:</strong> {exam.get('class_level', '')} | <strong>Time:</strong> {exam.get('time', '')}</p>
                    </div>
                """, unsafe_allow_html=True)
                
                if exam.get('section_a_mcqs'):
                    st.markdown('<div class="section-header">SECTION A – MCQs</div>', unsafe_allow_html=True)
                    for i, q in enumerate(exam['section_a_mcqs'], 1):
                        st.write(f"**Q{i}.** {q.get('question', '')}")
                        cols = st.columns(4)
                        for idx, opt in enumerate(q.get('options', [])):
                            cols[idx].text(opt)
                
                if exam.get('section_b_short'):
                    st.markdown('<div class="section-header">SECTION B – Short Questions</div>', unsafe_allow_html=True)
                    for i, q in enumerate(exam['section_b_short'], 1):
                        st.write(f"**Q{i}.** {q}")
                
                if exam.get('section_c_long'):
                    st.markdown('<div class="section-header">SECTION C – Long Questions</div>', unsafe_allow_html=True)
                    for i, q in enumerate(exam['section_c_long'], 1):
                        st.write(f"**Q{i}.** {q}")
                st.markdown("</div>", unsafe_allow_html=True)

            st.divider()
            
            col1, col2 = st.columns(2)
            with col1:
                st.markdown("### 📄 Question Paper (Student)")
                try:
                    doc_word = create_word_document(exam, show_answers=False, is_answer_key_only=False)
                    if doc_word:
                        st.download_button("Download Word", data=doc_word, file_name="Question_Paper.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                except Exception:
                    st.error("Could not generate Word file.")

                try:
                    doc_pdf = create_pdf_document(exam, show_answers=False, is_answer_key_only=False)
                    if doc_pdf:
                        st.download_button("Download PDF", data=doc_pdf, file_name="Question_Paper.pdf", mime="application/pdf")
                except Exception:
                    st.error("Could not generate PDF file.")

            with col2:
                st.markdown("### 🗝️ MCQ Answer Key (Teacher)")
                try:
                    key_word = create_word_document(exam, show_answers=True, is_answer_key_only=True)
                    if key_word:
                        st.download_button("Download Word", data=key_word, file_name="MCQ_Answer_Key.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")
                except Exception:
                    st.error("Could not generate Word file.")

                try:
                    key_pdf = create_pdf_document(exam, show_answers=True, is_answer_key_only=True)
                    if key_pdf:
                        st.download_button("Download PDF", data=key_pdf, file_name="MCQ_Answer_Key.pdf", mime="application/pdf")
                except Exception:
                    st.error("Could not generate PDF file.")

            if st.button("Start Over"):
                for key in list(st.session_state.keys()): del st.session_state[key]
                st.rerun()

if __name__ == "__main__":
    main()